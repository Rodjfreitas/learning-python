import pandas as pd
from docx import Document
from docx.shared import Inches

contagem = 1
confirmacao = input('Deseja iniciar:\n[1] - Sim\n[2] - Sair\n')

if int(confirmacao) == 1:

    # Importa os dados da planilha em formato de DataFrame
    df = pd.read_excel('RelatorioExemplos.xlsx')

    # Cria um novo arquivo word
    document = Document()

    # Adiciona título ao arquivo word
    document.add_heading('Relatório', 0)

    # Adiciona uma tabela com os dados da planilha ao arquivo word
    table = document.add_table(rows=len(df) + 1, cols=len(df.columns))
    for i in range(len(df.columns)):
        table.cell(0,i).text = df.columns[i]
    for i in range(len(df)):
        for j in range(len(df.columns)):
            table.cell(i + 1, j).text = str(df.values[i, j])

    # salva o arquivo do word
    document.save('{}{}{}'.format('arquivo_word_', contagem, '.docx'))
    contagem += 1
    print('{:=^50}'.format('Arquivo Gerado'))
else:
    print('{:=^50}'.format('Fim, Nada foi Executado.'))